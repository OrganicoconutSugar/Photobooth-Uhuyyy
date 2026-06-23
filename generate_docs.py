#!/usr/bin/env python3
"""
Generate comprehensive documentation for Photobooth project in .docx format.
Bahasa sehari-hari, lengkap dengan flowchart gambar dan penjelasan komponen.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os, textwrap, io

# ─────────────────────────────────────────────
# ─── HELPER: DRAW FLOWCHART ─────────────────
# ─────────────────────────────────────────────

FONT_PATH = "/System/Library/Fonts/Supplemental/Arial.ttf"

def _load_font(size):
    try:
        return ImageFont.truetype(FONT_PATH, size)
    except:
        return ImageFont.load_default()

def _draw_box(draw, x, y, w, h, text, fill="#F8F9F6", stroke="#9CAF88", font=None):
    """Draw rounded rectangle with text centered."""
    # Rectangle
    r = 12
    draw.rounded_rectangle([x, y, x+w, y+h], radius=r, fill=fill, outline=stroke, width=2)
    # Text
    if font is None:
        font = _load_font(14)
    # Word wrap
    lines = textwrap.wrap(text, width=20)
    total_h = len(lines) * 20
    start_y = y + (h - total_h) // 2
    for line in lines:
        bb = draw.textbbox((0, 0), line, font=font)
        tw = bb[2] - bb[0]
        tx = x + (w - tw) // 2
        draw.text((tx, start_y), line, fill="#2D3A2D", font=font)
        start_y += 22

def _draw_arrow(draw, x1, y1, x2, y2, color="#9CAF88"):
    """Draw a simple arrow line."""
    draw.line([(x1, y1), (x2, y2)], fill=color, width=2)
    # Arrow head at (x2, y2)
    dx = x2 - x1
    dy = y2 - y1
    length = (dx*dx + dy*dy) ** 0.5
    if length == 0:
        return
    ux, uy = dx/length, dy/length
    # Points for arrowhead
    ax1 = x2 - ux*10 - uy*6
    ay1 = y2 - uy*10 + ux*6
    ax2 = x2 - ux*10 + uy*6
    ay2 = y2 - uy*10 - ux*6
    draw.polygon([(x2, y2), (ax1, ay1), (ax2, ay2)], fill=color)

def create_flowchart_userflow():
    """User flow: how user navigates the app."""
    W, H = 600, 520
    img = Image.new("RGB", (W, H), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    font_big = _load_font(16)
    font_small = _load_font(12)
    bw = 140  # box width
    bh = 40   # box height

    title_font = _load_font(20)
    draw.text((W//2, 10), "Alur Pengguna (User Flow)", fill="#2D3A2D", font=title_font, anchor="mt")

    boxes = [
        (230, 50, "Landing Page", "#F0F7F0"),
        (230, 130, "Login / Register", "#E8F0E8"),
        (230, 210, "Kamera Capture", "#DDE8DD"),
        (80, 290, "Gallery\n(Lihat Foto)", "#E8F0E8"),
        (380, 290, "Templates\n(Pilih Frame)", "#E8F0E8"),
        (230, 370, "Download Strip", "#DDE8DD"),
        (230, 450, "Simpan ke Akun", "#C8DCC8"),
    ]
    for bx, by, text, bg in boxes:
        _draw_box(draw, bx, by, bw, bh, text, fill=bg)

    # Arrows
    _draw_arrow(draw, 300, 90, 300, 130)
    _draw_arrow(draw, 300, 170, 300, 210)
    _draw_arrow(draw, 230+bw//2, 250, 80+bw//2, 290)
    _draw_arrow(draw, 230+bw//2, 250, 380+bw//2, 290)
    _draw_arrow(draw, 80+bw//2, 330, 80+bw//2, 340)
    _draw_arrow(draw, 380+bw//2, 330, 380+bw//2, 340)
    # from gallery/templates back to center
    draw.line([(150, 340), (230, 340)], fill="#9CAF88", width=2)
    draw.line([(450, 340), (300, 340)], fill="#9CAF88", width=2)
    _draw_arrow(draw, 300, 340, 300, 370)

    _draw_arrow(draw, 300, 410, 300, 450)

    path = "/tmp/flow_userflow.png"
    img.save(path)
    return path

def create_flowchart_camera():
    """Camera capture flow."""
    W, H = 500, 550
    img = Image.new("RGB", (W, H), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    bw, bh = 140, 40

    title_font = _load_font(18)
    draw.text((W//2, 12), "Alur Kamera & Capture", fill="#2D3A2D", font=title_font, anchor="mt")

    boxes = [
        (180, 50, "Buka Halaman Kamera", "#F0F7F0"),
        (180, 120, "Cek Webcam", "#E8F0E8"),
        (180, 190, "Pilih Mode", "#DDE8DD"),
        (180, 260, "Tangkap 4 Foto", "#C8DCC8"),
        (180, 330, "Pilih Filter", "#DDE8DD"),
        (180, 400, "Download Strip", "#C8DCC8"),
        (180, 470, "Simpan ke Gallery", "#B8D0B8"),
    ]
    for bx, by, text, bg in boxes:
        _draw_box(draw, bx, by, bw, bh, text, fill=bg)

    # Vertical arrows
    for y in [90, 160, 230, 300, 370, 440]:
        _draw_arrow(draw, 250, y, 250, y+30)

    # Branch for webcam check
    draw.text((350, 130), "❌", fill="#E07070", font=_load_font(18))
    draw.text((50, 130), "✅", fill="#70B070", font=_load_font(18))
    # Lines for branch
    _draw_arrow(draw, 250, 160, 160, 160)
    _draw_arrow(draw, 160, 140, 160, 190)
    _draw_arrow(draw, 340, 140, 340, 190)
    _draw_arrow(draw, 340, 160, 250, 160)

    path = "/tmp/flow_camera.png"
    img.save(path)
    return path

def create_flowchart_auth():
    """Auth system flow."""
    W, H = 500, 520
    img = Image.new("RGB", (W, H), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    bw, bh = 160, 40

    title_font = _load_font(18)
    draw.text((W//2, 12), "Sistem Login & Auth", fill="#2D3A2D", font=title_font, anchor="mt")

    boxes = [
        (170, 50, "Buka Halaman Login", "#F0F7F0"),
        (170, 120, "Pilih Mode", "#E8F0E8"),
        (170, 190, "Isi Email + Password", "#DDE8DD"),
        (170, 260, "Kirim ke Server", "#C8DCC8"),
        (170, 330, "Cocokin Data?", "#B8D0B8"),
        (50, 410, "Berhasil!\n(Bikin Token JWT)", "#A0C0A0"),
        (290, 410, "Gagal\n(Muncul Error)", "#E8B0B0"),
    ]
    for bx, by, text, bg in boxes:
        _draw_box(draw, bx, by, bw, bh, text, fill=bg)

    # Vertical up to check
    _draw_arrow(draw, 250, 90, 250, 120)
    _draw_arrow(draw, 250, 160, 250, 190)
    _draw_arrow(draw, 250, 230, 250, 260)
    _draw_arrow(draw, 250, 300, 250, 330)
    # Branch
    _draw_arrow(draw, 250, 370, 130, 370)
    _draw_arrow(draw, 250, 370, 370, 370)
    _draw_arrow(draw, 130, 390, 130, 410)
    _draw_arrow(draw, 370, 390, 370, 410)

    path = "/tmp/flow_auth.png"
    img.save(path)
    return path

def create_flowchart_admin():
    """Admin flow."""
    W, H = 500, 420
    img = Image.new("RGB", (W, H), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    bw, bh = 160, 40

    title_font = _load_font(18)
    draw.text((W//2, 12), "Panel Admin", fill="#2D3A2D", font=title_font, anchor="mt")

    boxes = [
        (170, 60, "Login Admin", "#F0F7F0"),
        (170, 140, "Daftar Semua User", "#E8F0E8"),
        (50, 230, "Klik User\nLihat Gallery", "#DDE8DD"),
        (290, 230, "Hapus User\n(Perorangan)", "#E8B0B0"),
        (170, 310, "Hapus Inaktif /\nBersihkan Cache", "#D0C8D0"),
    ]
    for bx, by, text, bg in boxes:
        _draw_box(draw, bx, by, bw, bh, text, fill=bg)

    _draw_arrow(draw, 250, 100, 250, 140)
    _draw_arrow(draw, 250, 180, 130, 210)
    _draw_arrow(draw, 250, 180, 370, 210)
    _draw_arrow(draw, 250, 200, 250, 230)

    path = "/tmp/flow_admin.png"
    img.save(path)
    return path

# ─────────────────────────────────────────────
# ─── BUILD DOCX ──────────────────────────────
# ─────────────────────────────────────────────

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x2D, 0x3A, 0x2D)
    hs.font.name = 'Calibri'

# ─── COVER ───
for _ in range(4):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Photobooth - Dokumentasi Aplikasi")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x2D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Web App Photobooth Gratis berbasis React + Express")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x7A, 0x92, 0x76)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Dibuat oleh: Team's Uhuyyy Coperate\nTahun: 2026\nTema: Sage Maximalist")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x5A, 0x6E, 0x56)

doc.add_page_break()

# ─── DAFTAR ISI ───
doc.add_heading("Daftar Isi", level=1)
toc = [
    "1. Tentang Project Ini",
    "2. Tech Stack (Alat yang Dipakai)",
    "3. Alur User (User Flow)",
    "4. Alur Login & Auth",
    "5. Alur Kamera & Capture",
    "6. Alur Gallery",
    "7. Panel Admin",
    "8. Struktur File & Penjelasan",
    "  8.1 Bagian Server (Backend)",
    "  8.2 Bagian Frontend (Halaman)",
    "  8.3 Bagian Komponen",
    "  8.4 Bagian Library & Data",
    "9. Penjelasan Fitur Penting",
    "  9.1 AuthContext & JWT",
    "  9.2 Camera Capture & Filter",
    "  9.3 Smile Detection (MediaPipe)",
    "  9.4 Template Frame Styles",
    "  9.5 Gallery & Lightbox",
    "  9.6 Admin Panel",
    "10. Cara Jalanin Project",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ─── 1. TENTANG PROJECT ───
doc.add_heading("1. Tentang Project Ini", level=1)
doc.add_paragraph(
    "Jadi gini, aplikasi ini adalah web photobooth gratis. Lo bisa buka di browser, "
    "nyalain webcam, ambil foto 4 kali, kasih filter, trus download jadi 1 strip gambar "
    "atau simpen ke akun lo.\n\n"
    "Aplikasi ini punya fitur login, jadi foto-foto lo bisa disimpen di server dan "
    "dilihat lagi nanti. Ada juga panel admin buat ngelola user. Tema nya 'Sage Maximalist' "
    "— warnanya putih sama hijau sage (ijo-ijo soft gitu), ada border ornamental "
    "(hiasan sudut gitu), kesannya aesthetic maksimalis."
)

# ─── 2. TECH STACK ───
doc.add_heading("2. Tech Stack (Alat yang Dipakai)", level=1)

techs = [
    ("Frontend", "React 19, Vite 8, Tailwind CSS 3, React Router 6"),
    ("Backend", "Express 5 (pake ESM / import syntax)"),
    ("Database", "SQLite (pake better-sqlite3)"),
    ("Auth", "bcryptjs buat encrypt password, JWT buat token login"),
    ("Face Detection", "MediaPipe FaceMesh (dari CDN, bukan npm)"),
    ("Gambar", "Canvas API buat bikin strip foto, kompres avatar"),
    ("Lainnya", "ESLint buat ngelinting kode"),
]
for label, val in techs:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(val)

# ─── 3. USER FLOW ───
doc.add_heading("3. Alur User (User Flow)", level=1)
doc.add_paragraph(
    "Gambar di bawah ini nunjukin gimana user biasanya pakai aplikasi ini:"
)
img_path = create_flowchart_userflow()
doc.add_picture(img_path, width=Inches(4.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Penjelasan:\n"
    "1. User buka aplikasi dan sampai di Landing Page (halaman depan).\n"
    "2. Sebelum bisa foto, user harus login dulu. Kalo belum punya akun, bisa register.\n"
    "3. Setelah login, user bisa ke halaman Kamera buat foto.\n"
    "4. Sebelum foto, user bisa pilih template frame di halaman Templates.\n"
    "5. Abis 4 foto, user bisa download strip gambar atau simpan ke Gallery.\n"
    "6. Di Gallery bisa lihat foto lagi, download ulang, atau hapus."
)
os.remove(img_path)

# ─── 4. AUTH FLOW ───
doc.add_heading("4. Alur Login & Auth", level=1)
doc.add_paragraph(
    "Sistem auth pake JWT (token). Jadi pas login berhasil, server kasih token "
    "yang disimpen di localStorage browser. Token ini dipake tiap kali minta data "
    "dari server biar server tau siapa lo."
)
img_path = create_flowchart_auth()
doc.add_picture(img_path, width=Inches(4))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
os.remove(img_path)

doc.add_paragraph(
    "Halaman yang dilindungi (Kamera, Gallery, Templates) — kalo belum login, "
    "otomatis diarahkan ke halaman Login.\n\n"
    "Kalo lupa password: tinggal isi email, server bakal ngasih token reset "
    "(di demo langsung dikasih tokennya, nanti tinggal pake itu buat ganti password baru)."
)

# ─── 5. CAMERA FLOW ───
doc.add_heading("5. Alur Kamera & Capture", level=1)
doc.add_paragraph(
    "Fitur utama aplikasi. User bisa ambil 4 foto pake webcam.\n"
    "Ada 2 mode: Manual (pencet tombol sendiri) dan Senyum (otomatis pas senyum terdeteksi)."
)
img_path = create_flowchart_camera()
doc.add_picture(img_path, width=Inches(4))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
os.remove(img_path)

doc.add_paragraph(
    "Detail alur kamera:\n"
    "- Begitu masuk halaman, langsung cek webcam (minta izin browser).\n"
    "- Kalo ditolak, muncul mode simulasi pake gambar lokal.\n"
    "- User pilih mode Manual atau Senyum.\n"
    "- Mode Manual: tinggal pencet tombol, countdown 3 detik, jepret!\n"
    "- Mode Senyum: kamera deteksi mulut pake MediaPipe, kalo senyum terdeteksi 6x berturut-turut, otomatis njepret.\n"
    "- Abis 4 foto, tampil preview strip. User bisa download atau simpan ke gallery.\n"
    "- Ada 4 filter: Normal, Retro B&W, Warm Vintage, Neon Vibe."
)

# ─── 6. GALLERY ───
doc.add_heading("6. Alur Gallery", level=1)
doc.add_paragraph(
    "Gallery itu tempat liat foto-foto yang udah disimpan.\n\n"
    "- Kalo lagi login: data diambil dari server (API /api/sessions).\n"
    "- Kalo lagi ga login: gallery pake localStorage (data lokal).\n"
    "- Setiap session nampilin 4 thumbnail foto.\n"
    "- Klik thumbnail → muncul Lightbox (foto gede, bisa geser kanan/kiri, download per foto).\n"
    "- Tombol 'Unduh Strip' → regenerate ulang gambar strip + download.\n"
    "- Tombol hapus (ikon tempat sampah) → hapus session dari server.\n"
    "- Kalo gallery kosong, muncul pesan 'Belum Ada Cetakan'."
)

# ─── 7. ADMIN ───
doc.add_heading("7. Panel Admin", level=1)
doc.add_paragraph(
    "Admin bisa login pake akun khusus (admin@photobooth.app).\n"
    "Panel admin punya beberapa fitur:"
)
img_path = create_flowchart_admin()
doc.add_picture(img_path, width=Inches(4))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
os.remove(img_path)

admin_features = [
    "Melihat daftar semua user (email, username, role, jumlah foto, kapan terakhir login).",
    "Klik salah satu user → lihat foto-foto user itu.",
    "Tombol Hapus per user (tapi ga bisa hapus diri sendiri).",
    "Fitur 'Hapus Akun Tidak Aktif' — hapus user yang udah X hari ga login.",
    "Fitur 'Hapus Cache' — bersihin token reset password yang expired + VACUUM database.",
    "Sidebar: Admin ga liat menu Kamera (ga bisa akses kamera).",
]
for f in admin_features:
    doc.add_paragraph(f, style='List Bullet')

# ─── 8. STRUKTUR FILE ───
doc.add_heading("8. Struktur File & Penjelasan", level=1)

doc.add_heading("8.1 Bagian Server (Backend)", level=2)
doc.add_paragraph(
    "Server pake Express di port 3001. Ini dia file-file-nya:"
)
server_files = [
    ("server/index.js", "Entry point. Inisialisasi Express, mount routes, serve frontend (kalo udah dibuild), "
     "sama SPA fallback (biar React Router jalan). Ada seed admin otomatis pas startup."),
    ("server/db.js", "Setup database SQLite. Bikin tabel users, sessions, password_resets. "
     "Pake WAL mode biar cepet. Ada migration safe buat kolom baru (last_login, avatar)."),
    ("server/auth.js", "Semua soal login & register. Ada: POST /register, /login, GET /me, "
     "PUT /profile, POST /forgot-password, /reset-password. Juga middleware authMiddleware "
     "(cek token) sama adminMiddleware (cek role admin)."),
    ("server/sessions.js", "CRUD sessions. GET (ambil semua session user), POST (simpen session baru), "
     "DELETE (hapus session tertentu). Kalo user biasa, cuma bisa hapus punya sendiri. Admin bisa hapus siapa aja."),
    ("server/admin.js", "Endpoint khusus admin: GET /users (daftar user), GET /users/:id/sessions "
     "(foto user), POST /cleanup (bersihin cache), DELETE /users/:id (hapus user), "
     "DELETE /users/inactive/:days (hapus user ga aktif)."),
    ("server/seed.js", "Bikin user admin default pas pertama kali jalan (admin@photobooth.app / admin123)."),
]
for fname, desc in server_files:
    p = doc.add_paragraph()
    run = p.add_run(f"{fname}  —  ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

doc.add_heading("8.2 Bagian Frontend (Halaman)", level=2)
doc.add_paragraph(
    "Semua halaman ada di folder src/pages/. Masing-masing halaman diatur pake React Router."
)
pages = [
    ("Beranda.jsx", "Halaman depan aplikasi. Ada hero yang bagus pake gambar daun SVG, "
     "strip foto muter-muter, dan tombol 'Mulai Berfoto'."),
    ("Kamera.jsx", "Halaman utama buat foto. Pake webcam (getUserMedia), ada mode Manual & Senyum "
     "(MediaPipe FaceMesh), 4 filter, countdown 3 detik, flash overlay, preview 4 slot, "
     "tombol download strip, dan simpan ke gallery."),
    ("Gallery.jsx", "Tempat liat session yang udah disimpen. Bisa klik foto buat liat di Lightbox, "
     "download ulang strip, atau hapus. Fallback ke localStorage kalo ga login."),
    ("Templates.jsx", "Pilih frame template sebelum foto. Ada 8 gaya: Klasik, Polaroid, Vintage, "
     "Modern, Botanical, Double Frame, Cutie Cat, Retro Pop. Masing-masing punya warna sendiri."),
    ("Login.jsx", "Halaman login dengan 3 mode: Masuk, Buat Akun, Lupa Password. "
     "Pake ornament border dan SVG daun sebagai dekorasi."),
    ("Account.jsx", "Edit profil user. Bisa ganti username, bio, dan upload foto avatar. "
     "Email read-only. Ada tombol Keluar."),
    ("Admin.jsx", "Panel admin. Tabel daftar user, klik user lihat fotonya, "
     "tombol hapus, fitur cleanup, dan hapus user inaktif."),
]
for fname, desc in pages:
    p = doc.add_paragraph()
    run = p.add_run(f"{fname}  —  ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

doc.add_heading("8.3 Bagian Komponen", level=2)
components = [
    ("Sidebar.jsx", "Navbar glassmorphism di atas. Ada logo Photobooth, link Home/Kamera/Gallery/Templates, "
     "tombol Tentang & Privasi (modal), sama AccountMenu. Di mobile jadi hamburger menu."),
    ("AccountMenu.jsx", "Dropdown kecil di sidebar. Munculin avatar user, nama, sama link ke Account & Admin "
     "(kalo admin) plus tombol Keluar."),
    ("ProtectedRoute.jsx", "Bungkus halaman yang cuma bisa diakses kalo udah login. Kalo belum, "
     "redirect ke /login."),
    ("Lightbox.jsx", "Viewer foto full-screen. Bisa navigasi panah kiri/kanan, keyboard (Escape buat tutup), "
     "download per foto, dan indikator dot."),
]
for fname, desc in components:
    p = doc.add_paragraph()
    run = p.add_run(f"{fname}  —  ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

doc.add_heading("8.4 Bagian Library & Data", level=2)
libs = [
    ("context/AuthContext.jsx", "Jantung auth. Pake React Context buat nyimpen state user & token. "
     "Ada helper api() dan authedApi() buat fetch data dari server. Fungsi: login, register, logout, updateProfile. "
     "Pas pertama load, cek token di localStorage, kalo ada langsung verify ke server."),
    ("lib/galleryStore.js", "LocalStorage fallback buat nyimpen session. Key-nya 'photobooth.sessions.v1', "
     "max 10 session. Fungsi: loadSessions, saveSession, deleteSession, clearAllSessions."),
    ("lib/frameStyles.js", "Konfigurasi 8 template frame. Masing-masing punya id, name, warna (bg, border, accent, "
     "text), watermark, dan class CSS buat preview & kamera. Ada helper getFrameStyle()."),
]
for fname, desc in libs:
    p = doc.add_paragraph()
    run = p.add_run(f"{fname}  —  ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

# ─── 9. FITUR PENTING ───
doc.add_heading("9. Penjelasan Fitur Penting", level=1)

doc.add_heading("9.1 AuthContext & JWT", level=2)
doc.add_paragraph(
    "AuthContext itu kaya 'otak' dari sistem login. Cara kerjanya:\n\n"
    "- Pas user login/register, server ngirim JWT token (masa berlaku 7 hari).\n"
    "- Token ini disimpen di localStorage pake kunci 'pb_token'.\n"
    "- AuthContext punya helper `api()` yang otomatis nambahin header Authorization "
    "sama Content-Type tiap kali fetch data.\n"
    "- Helper `authedApi()` khusus buat endpoint yang butuh auth — tinggal kasih token.\n"
    "- Pas pertama halaman dibuka, AuthContext ngecek: 'Eh, ada token ga di localStorage?' "
    "Kalo ada, langsung verify ke /api/auth/me. Kalo valid, state user keisi. Kalo ga valid, "
    "token dihapus.\n"
    "- `ProtectedRoute` pake ini — 'Si user udah login belom? Kalo belom, lempar ke /login.'"
)

doc.add_heading("9.2 Camera Capture & Filter", level=2)
doc.add_paragraph(
    "Fitur kamera ada di Kamera.jsx. Gini cara kerjanya:\n\n"
    "Pertama, pas halaman dibuka, langsung coba ambil webcam pake `navigator.mediaDevices.getUserMedia()`.\n"
    "Kalo berhasil → stream video ditampilin di elemen <video>.\n"
    "Kalo ditolak → muncul layar 'Kamera Simulasi Aktif' pake gambar dari folder assets.\n\n"
    "Pas user pencet tombol capture (atau senyum terdeteksi):\n"
    "1. Hitung mundur 3 detik (countdown).\n"
    "2. Flash putih kedip.\n"
    "3. Gambar dari <video> di-capture ke <canvas>.\n"
    "4. Filter diterapkan (4 pilihan: normal, retro-bw, warm-vintage, neon-vibe).\n"
    "5. Hasilnya jadi data URL (base64) dan disimpen ke slot berikutnya.\n"
    "6. Ada 4 slot — masing-masing slot nampilin preview kecil.\n\n"
    "Abis 4 foto, muncul tombol 'Unduh Cetakan Foto (.png)':\n"
    "- Bikin canvas baru ukuran 400x1200 pixel.\n"
    "- 4 foto disusun vertikal dengan padding.\n"
    "- Ditambah border, watermark, dekorasi sesuai template.\n"
    "- Hasilnya di-download sebagai PNG.\n\n"
    "Tombol 'Simpan ke Gallery':\n"
    "- Kirim 4 foto (base64) ke server via POST /api/sessions.\n"
    "- Server simpen di tabel sessions, user_id sesuai token.\n"
    "- Muncul toast notifikasi sukses/gagal."
)

doc.add_heading("9.3 Smile Detection (MediaPipe)", level=2)
doc.add_paragraph(
    "Fitur deteksi senyum pake MediaPipe FaceMesh. Library-nya di-load dari CDN "
    "(bukan dari npm biar ga ribet bundling).\n\n"
    "Cara kerjanya:\n"
    "- Pas mode 'Senyum' dipilih, MediaPipe FaceMesh diinisialisasi.\n"
    "- Setiap frame video dianalisis — wajah di-scan, dicari 468 titik landmark.\n"
    "- Dari titik landmark itu, dihitung rasio tinggi mulut vs tinggi wajah.\n"
    "- Kalo rasionya >= 0.10 (artinya bukaan mulut cukup lebar = senyum), "
    "counter naik. Kalo 6 frame berturut-turut mendeteksi senyum → auto capture.\n"
    "- Ada progress bar biar user tau seberapa 'dekat' mood senyumnya.\n"
    "- Progress dihitung: minimal 100, ratio * 500.\n\n"
    "PENTING: Biar ga error pas countdown, state-nya disimpen di ref (bukan state biasa). "
    "Ref di-sync tiap render pake useEffect biar selalu up-to-date."
)

doc.add_heading("9.4 Template Frame Styles", level=2)
doc.add_paragraph(
    "Template frame ada 8 macam, didefinisikan di src/lib/frameStyles.js.\n"
    "Masing-masing template punya properti:\n\n"
    "- bg: warna background strip\n"
    "- border: warna border\n"
    "- accent: warna aksen (watermark, divider)\n"
    "- text: warna teks\n"
    "- watermark: teks watermark di strip\n"
    "- previewBg, previewBorder: class CSS buat preview card\n"
    "- slotBg, slotBorder: class CSS buat slot foto di preview\n"
    "- cameraBorder: class CSS buat border overlay di video kamera\n\n"
    "8 Template:\n"
    "1. Klasik — putih bersih, ijo sage\n"
    "2. Polaroid — border putih tebal, vintage\n"
    "3. Vintage — krem, sudut ornamental\n"
    "4. Modern Minimal — simpel, garis tipis\n"
    "5. Botanical — dedaunan, aksen gold\n"
    "6. Double Frame — garis ganda, stone tone\n"
    "7. Cutie Cat — pink gemes, watermark 'Meow!'\n"
    "8. Retro Pop — kuning neon, merah, watermark 'Pop!'\n\n"
    "Template dipilih dari halaman Templates, terus dikirim ke Kamera via URL parameter (?template=id). "
    "Kalo ga milih, default pake Klasik."
)

doc.add_heading("9.5 Gallery & Lightbox", level=2)
doc.add_paragraph(
    "Gallery:\n"
    "- Kalo user login: ambil data dari server (GET /api/sessions).\n"
    "- Kalo ga login: baca dari localStorage (galleryStore.js).\n"
    "- Tampilin grid kartu, masing-masing berisi 4 thumbnail foto.\n"
    "- Ada info template, tanggal, tombol Unduh Strip & Hapus.\n\n"
    "Lightbox:\n"
    "- Muncul pas thumbnail diklik.\n"
    "- Foto ditampilin gede di tengah, background gelap blur.\n"
    "- Navigasi: panah kiri/kanan, keyboard Arrow kiri/kanan, Escape buat tutup.\n"
    "- Indikator dot di bawah: klik langsung lompat ke foto itu.\n"
    "- Tombol 'Download Foto': download gambar yang sedang dilihat.\n"
    "- Scroll body dikunci (overflow: hidden) biar ga geser aneh."
)

doc.add_heading("9.6 Admin Panel", level=2)
doc.add_paragraph(
    "Panel admin di halaman /admin. Cuma bisa diakses user dengan role 'admin'.\n\n"
    "Tabel user:\n"
    "- Nampilin avatar, username, email, role, jumlah foto, tanggal buat, terakhir login.\n"
    "- Klik baris → lihat gallery foto user itu.\n"
    "- Tombol hapus (ikon sampah) → hapus user + semua foto-fotonya (CASCADE).\n\n"
    "Fitur lain:\n"
    "- 'Hapus Cache': bersihin token reset password yang expired + VACUUM database "
    "(biar ukuran file DB ga gede2 amat).\n"
    "- 'Hapus Akun Tidak Aktif': pilih jumlah hari (misal 30), trus server bakal "
    "hapus user yang login terakhir lebih dari X hari lalu. Admin ga bisa kehapus.\n"
    "- Admin gabisa lihat menu Kamera di sidebar (di-filter otomatis).\n"
    "- Kalo admin coba buka /kamera langsung, diarahkan ke /admin."
)

# ─── 10. CARA JALANIN ───
doc.add_heading("10. Cara Jalanin Project", level=1)
doc.add_paragraph(
    "Gampang, tinggal ikutin langkah-langkah ini:\n\n"
    "1. Install dependencies:\n"
    "   npm install\n\n"
    "2. Jalanin mode development (Vite + Express barengan):\n"
    "   npm run dev\n\n"
    "3. Buka browser:\n"
    "   Frontend: http://localhost:5173\n"
    "   API: http://localhost:3001\n\n"
    "4. Login admin:\n"
    "   Email: admin@photobooth.app\n"
    "   Password: admin123\n\n"
    "5. Build production:\n"
    "   npm run build\n\n"
    "6. Preview build:\n"
    "   npm run preview\n\n"
    "7. Linting:\n"
    "   npm run lint\n\n"
    "CATATAN: Kalo ada perubahan di file server, harus di-restart (Ctrl+C trus npm run dev lagi). "
    "Vite proxy udah diatur biar /api/* langsung ke Express."
)

# ─── SAVE ───
output_path = os.path.expanduser("~/Desktop/Dokumentasi_Photobooth.docx")
doc.save(output_path)
print(f"Dokumen berhasil dibuat: {output_path}")
